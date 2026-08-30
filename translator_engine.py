"""
translator_engine.py
Moteur LLM pour DOCX, XLSX, XLIFF
Support multi-format, reconstruction de balises, exports propres
"""
from platform import system
import time
import re
import csv
from urllib import response
import xml.etree.ElementTree as ET
from lxml import etree
from pathlib import Path
import logging
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger("aitranslator.engine")

SUPPORTED = {'.docx', '.xlsx', '.xliff', '.xlf'}
NS_X = 'urn:oasis:names:tc:xliff:document:1.2'
ET.register_namespace('', NS_X)

#── Segmentation multilingue ───────────────────────────────────────────────────────

def segment_text(text, source_lang="", use_segmentation=True):
    
    import pysbd

    if not use_segmentation or not text or not text.strip():
        return [text] if text else []

    lang = source_lang[:2].lower() if source_lang else "en"
    
    try:
        seg = pysbd.Segmenter(language=lang, clean=False)
        sents = seg.segment(text)
        sents = [s.strip() for s in sents if s.strip()]
        if sents:
            return sents
    except Exception:
        logger.warning("pysbd segmentation failed, fallback to regex")
        pass

    # Fallback léger : regex occidentale standard
    pattern = r"(?<=[.!?])\s+(?=[A-ZÀ-ÖØ-öø-ÿ])"
    sents = re.split(pattern, text)
    sents = [s.strip() for s in sents if s.strip()]
    return sents if sents else [text]


# ── Exceptions ───────────────────────────────────────────────────────────────
class TranslationStoppedException(Exception):
    """Exception levée quand l'utilisateur demande l'arrêt"""
    pass

# ── Routeur LLM ────────────────────────────────────────────────────────────
def translate_with_llm(text, source_lang, target_lang, context, 
                       style_instructions='', model_str="", api_key='', system_prompt_override=None):

    if not text or not text.strip():
        return text
    if isinstance(text, bytes):
        text = text.decode('utf-8', errors='ignore')
    text = text.replace('\x00', '')

    if '/' not in model_str:
        provider, model = model_str, 'default'
    else:
        provider, model = model_str.split("/", 1)
    
    provider = provider.strip()
    allowed = ("google", "openai", "qwen", "deepseek", "anthropic", "mistral", "ollama_cloud", "ollama_local")
    if provider not in allowed:
        #logger.error("Unknown provider: %s", provider)
        raise ValueError(f"Unknown provider: {provider}")

    style_block = f"\nSTYLE / TONE:\n{style_instructions.strip()}\n" if style_instructions and style_instructions.strip() else ""
    ctx_block = ""
    if context:
        ctx_block = "Recent context:\n" + "\n".join(f"- {c}" for c in context) + "\n"
    # Système personnalisé ou par défaut
    system_prompt = system_prompt_override if system_prompt_override is not None else (
        f"Professional translator.\n"
        f"ABSOLUTE RULES:\n"
        f"- Return ONLY the translation, nothing else\n"
        f"- NEVER repeat the source text in your response\n"
        f"- Do not include quotation marks, prefixes, or explanations\n"
        f"- Preserves the placeholders {{0}}, {{1}} exactly.\n"
    )

    user_prompt = (
        f"Translate from {source_lang} to {target_lang}.\n"
        f"Style guides:{style_block}\n"
        f"Context: {ctx_block}\n"
        f"Text to be translated:\n{text}"
    )
    if provider == "google":
        from google import genai
        if not api_key:
            logger.error("No API key provided.")
            raise ValueError("Google GenAI requires an API key. Please provide it.")
        if not model or model == "default":
            logger.error("No model provided.")
            raise ValueError(f"Missing {provider} model. Expected format: {provider}/<model>")
        client = genai.Client(api_key=api_key)
        config = {
            "system_instruction": system_prompt,
            "max_output_tokens": 4000
        }
        r = client.models.generate_content(model=model, contents=user_prompt, config=config)
        return r.text.strip()

    elif provider in ("openai", "qwen", "deepseek"):
        from openai import OpenAI
        if not api_key:
            logger.error("No API key provided.")
            raise ValueError(f"{provider} requires an API key. Please provide it.")
        if not model or model == "default":
            logger.error("No model provided.")
            raise ValueError(f"Missing {provider} model. Expected format: {provider}/<model>")
        urls = {
            "openai": "https://api.openai.com/v1",
            "qwen": "https://dashscope-intl.aliyuncs.com/compatible-mode/v1",
            "deepseek": "https://api.deepseek.com/v1"
        }
        client = OpenAI(api_key=api_key, base_url=urls[provider])
        r = client.chat.completions.create(
            model=model, messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.2, max_tokens=4000
        )
        return r.choices[0].message.content.strip()

    elif provider == "anthropic":
        import anthropic
        if not api_key:
            logger.error("No API key provided.")
            raise ValueError("Anthropic requires an API key. Please provide it.")
        if not model or model == "default":
            logger.error("Missing model. Expected format: provider/<model>")
            raise ValueError(f"Missing {provider} model. Expected format: {provider}/<model>")        
        client = anthropic.Anthropic(api_key=api_key)
        r = client.messages.create(
            model=model, max_tokens=4000,
            messages=[{"role": "user", "content": user_prompt}],
            system=system_prompt
        )
        return r.content[0].text.strip()
    elif provider == "mistral":
        from mistralai.client import Mistral
        if not api_key:
            logger.error("No API key provided.")
            raise ValueError("Mistral requires an API key. Please provide it.")
        if not model or model == "default":
            logger.error("Missing model. Expected format: provider/<model>")
            raise ValueError(f"Missing {provider} model. Expected format: {provider}/<model>")        
        client = Mistral(api_key=api_key)
        r = client.chat.complete(
                model=model, messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt},
                ],
                temperature=0.2, max_tokens=4000, response_format={"type": "text"}
            )
        return r.choices[0].message["content"].strip()
    elif provider == "mistral_v1":
        import requests
        import json
        if not api_key:
            logger.error("Mistral requires an API key. Please provide it.")
            raise ValueError(f"Mistral requires an API key. Expected format: {provider}/<model>")
        if not model or model == "default":
            logger.error("Missing model. Expected format: provider/<model>")
            raise ValueError(f"Missing Mistral model. Expected format: {provider}/<model>")
        url = "https://api.mistral.ai/v1/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}

        payload = {
            "model": model, 
            "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}],
            "stream": False,
            "temperature": 0.1,
            "max_tokens": 4000
        }
                
        try:
            response = requests.post(url, headers=headers, json=payload, timeout=120)
            response.raise_for_status()
            result = response.json()
            return result["choices"][0]["message"]["content"].strip()
        except requests.exceptions.ConnectionError:
            logger.error("API connection error.")
            raise Exception(f"Mistral is not running. Please start {provider} first.")
        except Exception as e:
            logger.error(f"Provider error: {e}")
            raise Exception(f"Provider error: {e}")

    elif provider == "ollama_cloud":
        import requests
        import os
        from ollama import Client
        if not api_key:
            logger.error("No API key provided.")
            raise ValueError("Ollama Cloud requires an API key. Please provide it.")
        if not model or model == "default":
            logger.error("Missing model. Expected format: provider/<model>")
            raise ValueError(f"Missing {provider} model. Expected format: {provider}/<model>") 
        ollama_cloud_host = os.getenv("OLLAMA_CLOUD_HOST", "https://ollama.com")
        try:
            client = Client(host=ollama_cloud_host, headers={"Authorization": f"Bearer {api_key}"}, timeout=120)
            r = client.chat(
                model=model,
                messages=[
                    {"role": "system", "content": system_prompt},
                    {"role": "user", "content": user_prompt}
                ]
            )
            return r["message"]["content"].strip()

        except Exception as e:
            logger.error(f"Ollama Cloud error: {e}")
            raise Exception(f"Ollama Cloud error: {e}")
    
    elif provider == "ollama_local":
        from ollama import Client
        import requests
        if not model or model == 'default':
            logger.error("Missing model. Expected format: provider/<model>")
            raise ValueError(f"Missing {provider} model. Expected format: {provider}/<model>, "
                              f"for example {provider}/llama3:latest")
        try:
            # Vérifier si le serveur Ollama local est en cours d'exécution
            response = requests.get("http://localhost:11434/api/tags", timeout=5)
            response.raise_for_status()
        except requests.exceptions.RequestException as e:
            logger.error("Ollama local is not running.")
            raise Exception(f"Ollama local is not running. Please start Ollama first: {e}")
        client = Client(host="http://localhost:11434")
        r = client.chat(
            model=model, 
            messages=[{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}]
        )
        return r["message"]["content"].strip()

    #logger.error("Unsupported provider: %s", provider)
    raise ValueError(f"Unsupported provider: {provider}")

# ── Utilitaires XLIFF ──────────────────────────────────────────────────────
def _extract(el):
    """
    Extrait le texte ET les balises internes d'un élément XLIFF.
    Protège les balises avec des placeholders {n} tout en conservant leur contenu.
    """
    tags = []  # Liste des balises protégées
    parts = []  # Liste des parties du texte
    tag_counter = 0  # Compteur pour les placeholders
    
    def walk(e, is_protected=False):
        nonlocal tag_counter
        
        # Texte avant les enfants
        if e.text:
            parts.append(e.text)
        
        for ch in e:
            loc = ch.tag.split('}')[-1] if '}' in ch.tag else ch.tag
            
            # Vérifier si c'est une balise à protéger
            if loc in ('g', 'x', 'bpt', 'ept', 'ph', 'it', 'mrk', 'sub'):
                # Protéger cette balise
                tag_index = tag_counter
                tag_counter += 1
                tags.append(ch)
                
                # Ajouter le placeholder d'ouverture
                parts.append(f'{{{tag_index}}}')
                
                # Traduire le contenu de la balise
                if ch.text:
                    parts.append(ch.text)
                
                # Traiter les enfants de la balise
                for child in ch:
                    walk(child, is_protected=True)
                
                # Ajouter le texte après les enfants (tail)
                if ch.tail:
                    parts.append(ch.tail)
                
                # Ajouter le placeholder de fermeture (optionnel)
                # Note: Certains formats n'ont pas de balise de fermeture explicite
                # Dans ce cas, on peut ignorer cette partie
                if loc in ('g', 'bpt', 'ph', 'mrk'):
                    # Ces balises ont généralement une fermeture
                    pass
                    
            else:
                # Balise non protégée - continuer la récursion
                walk(ch)
                if ch.tail:
                    parts.append(ch.tail)
    
    walk(el)
    
    # Filtrer les None et les chaînes vides
    parts = [p for p in parts if p is not None and p != '']
    
    return ''.join(parts), tags

def _rebuild(el, tr, tags):

    """
    Reconstruit un élément XLIFF avec ses balises internes.
    
    Args:
        el: Élément XML parent
        tr: Texte traduit avec placeholders {0}, {1}, ...
        tags: Liste des balises internes à réinsérer
    """

    import copy
    # Supprimer tous les enfants existants
    for child in list(el): 
        el.remove(child)
    el.text = None 
    last = None # Dernier élément ajouté (pour gérer les .tail)
    # Séparer le texte en parties (placeholders et texte)
    for part in re.split(r'({\d+})', tr):
        if not part:
            continue
        m = re.fullmatch(r'\{(\d+)\}', part)
        if m:
            # C'est un placeholder {n}
            idx = int(m.group(1))

            if idx < len(tags):
                # Insérer la balise correspondante
                tag = copy.deepcopy(tags[idx])
                tag.tail = ''
                el.append(tag)
                last = tag
            else:
                # Index invalide - traiter comme texte normal
                if last is not None:
                    last.tail = (last.tail or '') + part
                else:
                    el.text = (el.text or '') + part
        else:
            # C'est du texte normal
            if last is not None:
                last.tail = (last.tail or '') + part
            else:
                el.text = (el.text or '') + part
    #logger.debug(f"Rebuild: tr='{tr}', tags={len(tags)}")

def _clean_translation(translation, source_text):
    """
    Si la traduction contient le texte source colle, recupere uniquement la cible.
    Gere les variantes : source+cible, cible+source, separes par . / \n / espace.
    """
    if not translation or not source_text:
        return translation

    tr = translation.strip()
    src = source_text.strip()

    # Cas 1 : correspondance exacte — source presente dans la traduction
    if src in tr:
        # Source au debut : "source. traduction"
        after = tr.split(src, 1)[1].strip().lstrip('.,;:\n ')
        if after:
            return after
        # Source a la fin : "traduction. source"
        before = tr.split(src, 1)[0].strip().rstrip('.,;:\n ')
        if before:
            return before

    # Cas 2 : source tronquee — comparer les N premiers caracteres
    # Utile si le LLM a legerement reformate la source
    preview_len = min(40, len(src))
    src_preview = src[:preview_len]
    if len(src_preview) > 10 and tr.startswith(src_preview):
        # Trouver la fin approximative de la source dans la traduction
        # en cherchant le premier separateur de phrase apres la preview
        rest = tr[preview_len:]
        match = re.search(r'[.!?\n]', rest)
        if match:
            candidate = rest[match.start() + 1:].strip()
            if len(candidate) > 5:
                return candidate

    # Cas 3 : separation par saut de ligne
    if '\n' in tr:
        lines = [l.strip() for l in tr.split('\n') if l.strip()]
        if len(lines) == 2:
            # Determiner quelle ligne est la traduction
            # La ligne qui ne ressemble pas a la source
            for line in lines:
                if src[:20] not in line and len(line) > 5:
                    return line

    # Aucun pattern detecte : retourner tel quel
    return tr

def _clean_bilingual_source(txt, source_lang, target_lang):
    """
    Detecte et nettoie les segments sources qui contiennent deja leur traduction
    collee (pattern courant dans les exports Smartcat, SDL, Memsource).

    Exemples de segments bilingues detectes :
      "Hello world. Bonjour le monde."
      "Warning: danger. Avertissement : danger."

    Heuristique : si le texte depasse 1.5x la longueur attendue ET contient
    des scripts differents (latin/CJK/arabe), on tronque a la premiere moitie.
    Pour les paires de meme script, on ne peut pas trancher de facon sure
    et on retourne le texte complet — le prompt LLM renfonce gere le reste.
    """
    txt = txt.strip()

    # Detection CJK dans un texte majoritairement latin (ou vice versa)
    cjk_chars = len(re.findall(r'[\u4e00-\u9fff\u3040-\u30ff\uac00-\ud7af]', txt))
    latin_chars = len(re.findall(r'[a-zA-ZÀ-ÿ]', txt))
    total_chars = len(txt.replace(' ', ''))

    if total_chars == 0:
        return txt

    cjk_ratio = cjk_chars / total_chars
    latin_ratio = latin_chars / total_chars

    # Si le texte est significativement mixte latin+CJK : c'est bilingue
    # Seuils : au moins 15% de chaque script
    if cjk_ratio > 0.15 and latin_ratio > 0.15:
        # Strategie : garder seulement la partie source
        # On cherche la frontiere entre les deux langues
        # Cas 1 : source=latin, cible=CJK -> garder la partie latine
        # Cas 2 : source=CJK, cible=latin -> garder la partie CJK

        # Detecter la langue source dominante
        src_is_latin = source_lang.startswith(('en', 'fr', 'de', 'es', 'pt', 'it', 'nl'))
        src_is_cjk = source_lang.startswith(('zh', 'ja', 'ko'))

        if src_is_latin:
            # Garder uniquement la partie avant le premier caractere CJK
            match = re.search(r'[\u4e00-\u9fff\u3040-\u30ff\uac00-\ud7af]', txt)
            if match:
                return txt[:match.start()].strip()
        elif src_is_cjk:
            # Garder uniquement la partie CJK (avant le latin)
            match = re.search(r'[a-zA-ZÀ-ÿ]{3,}', txt)
            if match:
                return txt[:match.start()].strip()
    return txt

def _validate_translation(source_txt, translation, source_lang, target_lang):
    """
    Valide que la traduction ne contient pas le texte source colle.
    Retourne (ok, raison).
    """
    if not translation or not translation.strip():
        return False, "No translation"

    # Verifier que la traduction n'est pas identique a la source
    if translation.strip() == source_txt.strip():
        return False, "Translation identical to source"

    # Verifier que la traduction ne commence pas par le texte source
    # (pattern : "source. traduction" ou "source traduction")
    src_preview = source_txt.strip()[:50]
    if len(src_preview) > 10 and translation.startswith(src_preview):
        return False, "Translation prefixed with source"

    # Verifier les placeholders : tous ceux de la source doivent etre dans la cible
    src_tokens = set(re.findall(r'\{\d+\}', source_txt))
    tr_tokens = set(re.findall(r'\{\d+\}', translation))
    missing = src_tokens - tr_tokens
    if missing:
        return False, f"Missing placeholders: {missing}"
    extra = tr_tokens - src_tokens
    if extra:
        return False, f"Extra placeholders: {extra}"
    return True, None

# ── Formats de sortie ──────────────────────────────────────────────────────
# XLSX bilingue simple
def generate_bilingual_xlsx(pairs, sl, tl, out_path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Bilingual"
    ws.append([f"Source ({sl})", f"Target ({tl})"])
    for s, t in pairs:
        ws.append([s, t])
    ws.column_dimensions['A'].width = 50
    ws.column_dimensions['B'].width = 50
    out = str(Path(out_path).with_suffix('.xlsx'))
    wb.save(out)
    return out

# TMX standard
def _clean_for_xml(text):
    """
    Supprime les caractères de contrôle invalides en XML.
    """
    if text is None:
        return ""

    text = str(text)

    # XML autorise \t, \n, \r mais pas la plupart des autres caractères de contrôle
    text = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', text)

    return text


def _normalize_lang_code(lang):
    """
    Retourne un code de langue utilisable dans xml:lang.
    Exemple: fr-FR -> fr-FR
             fr_FR -> fr-FR
             Français -> fr (si mapping externe non disponible, on garde tel quel)
    """
    if not lang:
        return "und"

    lang = str(lang).strip()

    if not lang:
        return "und"

    # Remplace les underscores par des tirets pour BCP47
    lang = lang.replace("_", "-")

    return lang

def generate_tmx(pairs, source_lang, target_lang, out_path):
    import xml.dom.minidom as minidom
    from datetime import datetime, timezone

    doc = minidom.Document()

    tmx = doc.createElement('tmx')
    tmx.setAttribute('version', '1.4')
    doc.appendChild(tmx)

    header = doc.createElement('header')
    header.setAttribute('creationtool', 'AI Translator')
    header.setAttribute('adminlang', 'en')
    header.setAttribute('srclang', _normalize_lang_code(source_lang))
    header.setAttribute('creationdate', datetime.now(timezone.utc).strftime('%Y%m%dT%H%M%SZ'))
    tmx.appendChild(header)

    body = doc.createElement('body')
    tmx.appendChild(body)

    for src, tgt in pairs:
        tu = doc.createElement('tu')

        for lang, txt in [(source_lang, src), (target_lang, tgt)]:
            tuv = doc.createElement('tuv')
            tuv.setAttribute('xml:lang', _normalize_lang_code(lang))

            seg = doc.createElement('seg')
            seg.appendChild(doc.createTextNode(_clean_for_xml(txt)))

            tuv.appendChild(seg)
            tu.appendChild(tuv)

        body.appendChild(tu)

    out = str(Path(out_path).with_suffix('.tmx'))

    with open(out, 'wb') as f:
        f.write(doc.toxml(encoding='utf-8'))

    return out

def generate_tsv(pairs, source_lang, target_lang, out_path):
    out = str(Path(out_path).with_suffix('.tsv'))

    source_header = f"Source ({source_lang})" if source_lang else "Source"
    target_header = f"Target ({target_lang})" if target_lang else "Target"

    with open(out, 'w', encoding='utf-8-sig', newline='') as f:
        writer = csv.writer(
            f,
            delimiter='\t',
            quoting=csv.QUOTE_MINIMAL
        )
        writer.writerow([source_header, target_header])

        for s, t in pairs:
            writer.writerow([
                _clean_for_tsv(s),
                _clean_for_tsv(t)
            ])

    return out

# XLIFF standard
def _xliff_tag(tag):
    return f'{{{NS_X}}}{tag}'

def generate_xliff(pairs, source_lang, target_lang, out_path, original_filename=""):
    """Générer xliff standard à partir du paire de langues"""
    import xml.etree.ElementTree as ET

    ET.register_namespace('', NS_X)

    xliff = ET.Element(_xliff_tag('xliff'), version="1.2")

    file_elem = ET.SubElement(xliff, _xliff_tag('file'), {
        'original': Path(original_filename).name if original_filename else "source",
        'source-language': source_lang,
        'target-language': target_lang,
        'datatype': 'plaintext'
    })

    body = ET.SubElement(file_elem, _xliff_tag('body'))

    for i, (src, tgt) in enumerate(pairs, 1):
        tu = ET.SubElement(body, _xliff_tag('trans-unit'), id=str(i))

        source_el = ET.SubElement(tu, _xliff_tag('source'))
        source_el.text = src

        target_el = ET.SubElement(tu, _xliff_tag('target'), {'state': 'needs-review-translation'})
        target_el.text = tgt

    out = str(Path(out_path).with_suffix('.xliff'))
    ET.ElementTree(xliff).write(out, encoding='utf-8', xml_declaration=True)
    return out

# ── DOCX ────────────────────────────────────────────────────────────────
# ── SDT : détection bloc vs inline ──────────────────────────────────────
# Vérifier si un élément est dans une balise SDT
def is_inside_sdt(element):
    """
    Vérifie si un élément XML est un descendant d'une balise w:sdt.
    Remonte l'arbre XML, nœud par nœud, pour vérifier si l'élément donné se trouve à l'intérieur d'une balise <w:sdt>.
    """
    parent = element.getparent()
    # Parcourir les parents successifs
    while parent is not None:
        if parent.tag == qn('w:sdt'): # Si le parent est une balise <w:sdt>, retourner True
            return True # Sortie anticipée / Succès. Stoppe immédiatement dès qu'une balise w:sdt est trouvée.
        # Si ce n'est pas le bon tag, parent = parent.getparent() passe au niveau supérieur 
        # (le grand-parent, puis le arrière-grand-parent, etc.).
        # La boucle s'arrête quand on atteint la racine de l'arbre, car le parent de la racine vaut None.
        parent = parent.getparent() 
    return False # Sortie par défaut / Échec
    # Placé en dehors de la boucle while, il ne s'exécute que si la boucle s'est terminée sans jamais avoir rencontré de balise w:sdt (c'est-à-dire une fois arrivé à la racine parent is None). Cela confirme que l'élément n'est pas dans un SDT.

def is_inline_sdt(sdt):
    """Vrai si ce <w:sdt> est un enfant DIRECT d'un <w:p> (content control
    inline, ex: tags goog_rdk_* des exports Google Docs, qui encapsulent
    un mot ou une portion de phrase à l'intérieur d'un paragraphe).
    Faux s'il est au niveau bloc (frère de <w:p> dans <w:body>/<w:tc>,
    et contenant lui-même des <w:p> dans son sdtContent)."""
    parent = sdt.getparent()
    return parent is not None and parent.tag == qn('w:p')

# ── Lecture/écriture UNIFIÉE d'un paragraphe (runs directs + SDT inline) ──
# CORRECTIF : p.text / p.runs de python-docx ne voient QUE les <w:r>
# enfants DIRECTS d'un <w:p> -- ils ignorent tout texte imbriqué dans un
# <w:sdt>. Sur un paragraphe mixte (runs + SDT inline goog_rdk_*), cela
# provoquait un double traitement : le paragraphe était traduit "troué"
# (mots des SDT manquants) et réinjecté dans runs[0], tandis que chaque
# SDT était traduit isolément (hors contexte) et réinjecté à sa place —
# cassant l'ordre de lecture (les fragments SDT semblaient "poussés" en
# fin de paragraphe une fois les runs directs environnants vidés).
#
# Les fonctions ci-dessous traitent tout le paragraphe comme UNE SEULE
# séquence ordonnée de runs, quelle que soit leur provenance.
# Extraction du texte SDT
NS_W = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
_NSMAP = {'w': NS_W}
def get_ordered_runs(p_element):
    """Liste ORDONNÉE des <w:r> d'un paragraphe, en incluant ceux
    imbriqués dans des <w:sdt> inline (y compris SDT empilés, ex:
    goog_rdk_0 > goog_rdk_1 > w:r), dans l'ordre réel du document."""
    runs = []
    for child in p_element:
        tag = etree.QName(child).localname
        if tag == 'r':
            runs.append(child)
        elif tag == 'sdt':
            runs.extend(child.findall('.//w:r', _NSMAP))
    return runs

def get_run_text(run):
    return ''.join(t.text or '' for t in run.findall(qn('w:t')))

def get_paragraph_full_text(p_element):
    """Remplace p.text : concatène le texte de TOUS les runs du
    paragraphe, y compris ceux imbriqués dans des SDT inline."""
    return ''.join(get_run_text(r) for r in get_ordered_runs(p_element))

def set_paragraph_text(p_element, new_text):
    """Remplace le pattern `element.runs[0].text = t` / vidage des runs
    suivants, mais sur la séquence UNIFIÉE (runs directs + runs dans SDT
    inline). Préserve le formatage du premier run porteur de texte, où
    qu'il soit dans l'arbre (direct ou dans un SDT)."""
    runs = get_ordered_runs(p_element)
    if not runs:
        return

    first = runs[0]
    t_elem = first.find(qn('w:t'))
    if t_elem is None:
        t_elem = OxmlElement('w:t') # Si aucun texte n'existe, on en crée un
        first.append(t_elem)
    t_elem.text = new_text # On remplace le texte du premier run
    t_elem.set(qn('xml:space'), 'preserve') # On force l'attribut xml:space="preserve" pour éviter que Word ne supprime les espaces en début/fin de texte
    # Vider d'éventuels w:t supplémentaires dans ce même run
    for extra_t in first.findall(qn('w:t'))[1:]:
        extra_t.text = ""

    # Vider tous les runs suivants, direct ou dans un SDT inline
    for r in runs[1:]:
        for t in r.findall(qn('w:t')):
            t.text = ""

# Extraction du texte SDT (BLOC uniquement — les SDT inline sont
# désormais couverts par get_paragraph_full_text / set_paragraph_text
# sur leur paragraphe parent, pour éviter le double traitement)
def extract_sdt_text(doc):
    """Extrait le texte des balises w:sdt de premier niveau, EN EXCLUANT
    les SDT inline (imbriqués dans un <w:p>). Ne restent que les SDT
    bloc (contenant leurs propres <w:p>, au niveau body/cell) — ex:
    champs de formulaire structurés."""
    
    sdt_elements = []

    # On cherche les w:sdt à la racine ou dans le corps du document
    for sdt in doc.element.findall('.//w:sdt', _NSMAP):
        # On s'assure de ne pas prendre les sous-sdt (sdt imbriqués)
        if is_inside_sdt(sdt):
            continue
        # On exclut les SDT inline (traités via la boucle de paragraphe)
        if is_inline_sdt(sdt):
            continue

        text_runs = sdt.findall('.//w:t', _NSMAP)
        text = ''.join([t.text for t in text_runs if t.text])
        if text.strip():
            sdt_elements.append((sdt, text.strip()))
    return sdt_elements

# Remplacement du texte SDT BLOC (Préserve le formatage)
# Utilisé uniquement pour les SDT de niveau bloc (contenant leurs
# propres <w:p>) — les SDT inline passent désormais par
# set_paragraph_text sur leur paragraphe parent.
def replace_sdt_text(sdt, new_text):
    """Remplace le texte d'un w:sdt BLOC en préservant le formatage du
    premier run. Conserve un filet de sécurité pour le cas inline
    (sdtContent sans <w:p>) au cas où un SDT isolé, non rattaché à un
    paragraphe, se présenterait malgré tout."""
    
    sdt_content = sdt.find(qn('w:sdtContent'))
    if sdt_content is None:
        return

    paragraphs = sdt_content.findall(qn('w:p'))

    if paragraphs:
        # 1. Insérer le nouveau texte dans le premier paragraphe
        first_p = paragraphs[0]
        runs = first_p.findall(qn('w:r'))

        if not runs:
            # Si aucun run n'existe, on en crée un
            r = OxmlElement('w:r')
            t_elem = OxmlElement('w:t')
            t_elem.text = new_text
            t_elem.set(qn('xml:space'), 'preserve')
            r.append(t_elem)
            first_p.append(r)
        else:
            # On utilise le premier run pour garder le formatage (gras, couleur, etc.)
            first_r = runs[0]
            t_elem = first_r.find(qn('w:t'))
            if t_elem is None:
                t_elem = OxmlElement('w:t')
                first_r.append(t_elem)

            t_elem.text = new_text
            t_elem.set(qn('xml:space'), 'preserve')

            # On vide les runs suivants du premier paragraphe
            for r in runs[1:]:
                for t in r.findall(qn('w:t')):
                    t.text = ""

        # 2. Vider le texte des paragraphes suivants dans la balise SDT
        for p in paragraphs[1:]:
            for r in p.findall(qn('w:r')):
                for t in r.findall(qn('w:t')):
                    t.text = ""
    else:
        # Filet de sécurité : sdtContent sans <w:p> (SDT isolé, non
        # rattaché à un paragraphe parent). Recherche en profondeur des
        # runs, remplacement du premier, vidage des suivants.
        runs = sdt_content.findall('.//w:r', _NSMAP)
        if not runs:
            return
        first_r = runs[0]
        t_elem = first_r.find(qn('w:t'))
        if t_elem is None:
            t_elem = OxmlElement('w:t')
            first_r.append(t_elem)
        t_elem.text = new_text
        t_elem.set(qn('xml:space'), 'preserve')
        for r in runs[1:]:
            for t in r.findall(qn('w:t')):
                t.text = ""

def translate_docx(inp, out, api_key, sl, tl, model, delay, style_instructions, 
                   output_formats, cb=None, stop_check=None, 
                   use_segmentation=False, context_size=3):
    """
    Traduit un fichier DOCX.
    
    Args:
        use_segmentation: Si True, segmente les paragraphes en phrases
        context_size: Nombre de phrases de contexte (0, 3, ou 5)
    """
    from docx import Document
    doc = Document(inp)
    segments = []
    # 1. Parser les paragraphes (EN IGNORANT ceux qui sont dans un SDT BLOC ; 
    #    les SDT inline sont désormais inclus automatiquement via get_paragraph_full_text)
    for p in doc.paragraphs:
        if is_inside_sdt(p._element):
            continue  # Paragraphe emprisonné dans un SDT bloc : sera traité par la logique SDT plus bas
        text = get_paragraph_full_text(p._element).strip()
        if text:
            if use_segmentation:
                sentences = segment_text(text, sl, use_segmentation=True) #Segmenter le paragraphe en phrases
                for sentence in sentences:
                    if sentence.strip():
                        segments.append(('p', p, sentence, True))#True = phrase segmentée
            else:
                segments.append(('p', p, text, False))#False = paragraphe entier

    # 2. Parser les tableaux (EN IGNORANT ceux qui sont dans un SDT bloc)
    for table in doc.tables:
        if is_inside_sdt(table._tbl):
            continue
            
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    text = get_paragraph_full_text(p._element).strip()
                    if text:
                        if use_segmentation:
                            sentences = segment_text(text, sl, use_segmentation=True)
                            for sentence in sentences:
                                if sentence.strip():
                                    segments.append(('cell', p, sentence, True))
                        else:
                            segments.append(('cell', p, text, False))
    # 3. Extraire les balises <w:sdt> BLOC uniquement (les SDT inline
    #    ont déjà été inclus dans le texte des paragraphes ci-dessus)
    sdt_data = extract_sdt_text(doc)
    for sdt, text in sdt_data:
        # On stocke l'élément sdt et le texte original
        if use_segmentation and len(text) > 100: # Seulement si le texte est long
            sentences = segment_text(text, sl, use_segmentation=True)
            for sentence in sentences:
                if sentence.strip():
                    segments.append(('sdt', sdt, sentence, True))
        else:
            segments.append(('sdt', sdt, text, False))

    total = len(segments)
    ctx = []
    pairs = []
    translated_buffer = {}   # Pour paragraphes et cellules (l'élément est toujours un Paragraph)
    sdt_buffer = {}          # Pour les <w:sdt>

    for i, (seg_type, element, original, is_sentence) in enumerate(segments):
        # Vérifier si l'arrêt est demandé
        if stop_check and stop_check():
            if cb: cb(total, total, "Translation stopped by the user.", True)
            raise TranslationStoppedException("Stop requested by the user")

        if cb:
            msg = f"Phrase {i+1}/{total}" if is_sentence else f"Segment {i+1}/{total}"
            cb(i, total, msg, False)

        try:
            # Construire le contexte selon le paramètre context_size
            context_for_llm = ctx[-context_size:] if context_size > 0 and len(ctx) >= context_size else ctx[:]
            t = translate_with_llm(original, sl, tl, context_for_llm, style_instructions, model, api_key)
            pairs.append((original, t))
            ctx.append(t)
            if len(ctx) > max(context_size, 3): # Limiter le contexte selon la taille choisie
                ctx.pop(0)

            # --- Gestion selon le type de segment ---
            if seg_type in ('p', 'cell'): # 'element' est TOUJOURS un objet Paragraph (même pour les cellules)
                if is_sentence: # Si c'est une phrase segmentée, on stocke la traduction
                    if element not in translated_buffer:
                        translated_buffer[element] = []
                    translated_buffer[element].append(t)
                else:
                    # Paragraphe entier : remplacer via la séquence unifiée
                    # (runs directs + runs imbriqués dans un SDT inline)
                    set_paragraph_text(element._element, t)

            elif seg_type == 'sdt': # Champs de texte structuré (SDT bloc) stocké dans sdt_buffer
                if is_sentence:
                    if element not in sdt_buffer:
                        sdt_buffer[element] = []
                    sdt_buffer[element].append(t)
                else:
                    replace_sdt_text(element, t) # Remplacer directement le texte du sdt

        except Exception as e:
            #logger.error("Error translating segment %s: %s", i, e)
            if cb: cb(i, total, f"Error {i+1}: {e}", False)
            
            # En cas d'erreur, garder l'original
            if is_sentence:
                buffer = sdt_buffer if seg_type == 'sdt' else translated_buffer
                if element not in buffer: buffer[element] = []
                buffer[element].append(original)
                # Pour les paragraphes/cellules, on pourrait ne rien faire ou laisser le texte original
            else:
                if seg_type == 'sdt': replace_sdt_text(element, original)
            pairs.append((original, original))

        if i < total - 1:
            time.sleep(delay)


    # --- RECONSTRUCTION DES SEGMENTS (Buffer) ---
    # Note: translated_buffer ne contient QUE des objets Paragraph (corps ou cellules)
    for element, translated_sentences in translated_buffer.items():
        full_translation = ' '.join(translated_sentences)
        set_paragraph_text(element._element, full_translation)
    # Reconstruire les <w:sdt> BLOC segmentés
    for sdt, translated_sentences in sdt_buffer.items():
        full_translation = ' '.join(translated_sentences)
        replace_sdt_text(sdt, full_translation)

    # Générer les fichiers de sortie
    paths = []
    for fmt in output_formats:
        #logger.info(f"Processing format: '{fmt}'")
        if fmt == 'original':
            out_path = str(Path(out).with_suffix('.docx'))
            doc.save(out_path)
            # Vérifier que le chemin est absolu
            #logger.info(f"DOCX saved to: {out_path}")
            #logger.info(f"File exists: {Path(out_path).exists()}")
            paths.append(out_path)
        elif fmt == 'bilingual_xlsx':
            paths.append(generate_bilingual_xlsx(pairs, sl, tl, out + '_bilingual'))
        elif fmt == 'tmx':
            paths.append(generate_tmx(pairs, sl, tl, out + '_bilingual'))
        elif fmt == 'tsv':
            paths.append(generate_tsv(pairs, sl, tl, out + '_bilingual'))
        elif fmt == 'xliff':
            paths.append(generate_xliff(pairs, sl, tl, out + '_bilingual', inp))
    
    if cb: cb(total, total, f"DOCX completed - {len(paths)} file(s) generated.", True)
    return paths

# ── XLSX ───────────────────────────────────────────────────────────────────
def translate_xlsx(inp, out, api_key, sl, tl, model, delay, style_instructions, 
                   output_formats, cb=None, stop_check=None,
                   use_segmentation=False, context_size=3):
    """
    Traduit un fichier XLSX.
    
    Args:
        use_segmentation: Si True, segmente les cellules en phrases
        context_size: Nombre de phrases de contexte (0, 3, ou 5)
    """
    import openpyxl
    wb = openpyxl.load_workbook(inp)
    
    cells = []
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and cell.value.strip():
                    text = cell.value.strip()
                    
                    if use_segmentation and len(text) > 100:
                        sentences = segment_text(text, sl, use_segmentation=True)
                        for sentence in sentences:
                            if sentence.strip():
                                cells.append((ws, cell, sentence, True))
                    else:
                        cells.append((ws, cell, text, False))
    
    total = len(cells)
    ctx = []
    pairs = []
    cell_buffer = {}  # Buffer pour reconstruire les cellules segmentées
    
    for i, (ws, cell, original, is_sentence) in enumerate(cells):
        # Vérifier si l'arrêt est demandé
        if stop_check and stop_check():
            if cb: 
                logger.exception("Stop Translation requested by the user")
                cb(total, total, "Translation stopped by the user", True)
            raise TranslationStoppedException("Stop requested by the user")
        
        if cb:
            msg = f"Phrase {i+1}/{total}" if is_sentence else f"Cellule {i+1}/{total}"
            cb(i, total, msg, False)
        
        try:
            # Construire le contexte
            context_for_llm = []
            if context_size > 0:
                context_for_llm = ctx[-context_size:] if len(ctx) >= context_size else ctx[:]
            
            t = translate_with_llm(original, sl, tl, context_for_llm, 
                                  style_instructions, model, api_key)
            
            pairs.append((original, t))
            ctx.append(t)
            
            if len(ctx) > max(context_size, 3):
                ctx.pop(0)
            
            if is_sentence:
                if cell not in cell_buffer:
                    cell_buffer[cell] = []
                cell_buffer[cell].append(t)
            else:
                cell.value = t
                
        except Exception as e:
            if cb: 
                #logger.error("Error translating segment %s", i)
                cb(i, total, f"Error {i+1}: {e}", False)
            if is_sentence:
                if cell not in cell_buffer:
                    cell_buffer[cell] = []
                cell_buffer[cell].append(original)
            pairs.append((original, original))
        
        if i < total - 1:
            time.sleep(delay)
    
    # Reconstruire les cellules segmentées
    for cell, translated_sentences in cell_buffer.items():
        cell.value = ' '.join(translated_sentences)
    
    # Générer les fichiers de sortie
    paths = []
    for fmt in output_formats:
        if fmt == 'original':
            out_path = str(Path(out).with_suffix('.xlsx'))
            wb.save(out_path)
            paths.append(out_path)
        elif fmt == 'bilingual_xlsx':
            paths.append(generate_bilingual_xlsx(pairs, sl, tl, out + '_bilingual'))
        elif fmt == 'tmx':
            paths.append(generate_tmx(pairs, sl, tl, out + '_bilingual'))
        elif fmt == 'tsv':
            paths.append(generate_tsv(pairs, sl, tl, out + '_bilingual'))
        elif fmt == 'xliff':
            paths.append(generate_xliff(pairs, sl, tl, out + '_bilingual', inp))
    
    if cb: cb(total, total, f"XLSX completed - {len(paths)} file(s) generated.", True)
    return paths

# ── XLIFF ──────────────────────────────────────────────────────────────────
def translate_xliff_file(inp, out, api_key, sl, tl, model, delay, skip_translated, style_instructions, output_formats, cb=None, stop_check=None, context_size=3):
    tree = ET.parse(inp)
    root = tree.getroot()
    units = root.findall(f'.//{{{NS_X}}}trans-unit')
    
    todo = []
    for u in units:
        src = u.find(f'{{{NS_X}}}source')
        if src is None: continue
        txt, tags = _extract(src)
        if not txt.strip(): continue
        if skip_translated:
            tgt = u.find(f'{{{NS_X}}}target')
            if tgt is not None and tgt.get('state', '') in ('translated', 'signed-off', 'final'):
                continue
        todo.append((u, src, txt, tags))

    total = len(todo)
    if total == 0:
        if cb: cb(0, 0, "No segments to translate.", True)
        return []

    ctx = []
    pairs = []
    for i, (u, src, txt, itags) in enumerate(todo):
        if stop_check and stop_check():
            if cb: 
                logger.exception("Stop Translation requested by the user")
                cb(total, total, "Translation stopped by the user.", True)
            raise TranslationStoppedException("Stop requested by the user")
        if cb: cb(i, total, f"Segment {i+1}/{total}", False)
        try:
            # Nettoyer les sources bilingues
            clean_txt = _clean_bilingual_source(txt, sl, tl)

            # Traduction avec retry si validation echoue
            tr = None
            for attempt in range(3):
                context_for_llm = ctx[-context_size:] if context_size > 0 else []
                candidate = translate_with_llm(
                    clean_txt, sl, tl, context_for_llm, style_instructions, model, api_key
                )
                ok, reason = _validate_translation(clean_txt, candidate, sl, tl)
                if ok:
                    # Nettoyage supplementaire par securite
                    tr = _clean_translation(candidate, clean_txt)
                    break                    

                # Si echec : renforcer le prompt au retry suivant
                # en ajoutant le probleme detecte dans le contexte
                if cb:
                    cb(i, total, f"Segment {i+1} retry {attempt+1} ({reason})", False)
                # Petit delai avant retry
                time.sleep(0.5)
            # Fallback si tous les retries echouent
            if tr is None:
                # Utiliser le candidat meme imparfait plutot que planter
                tr = candidate if candidate else clean_txt # On retourne la source si echec
                  
            if not tr.strip(): tr = clean_txt

            tgt = u.find(f'{{{NS_X}}}target')
            if tgt is None: tgt = ET.SubElement(u, f'{{{NS_X}}}target')
            tgt.set('state', 'needs-review-translation')
            _rebuild(tgt, tr, itags)
            pairs.append((clean_txt, tr))
            if context_size > 0:
                ctx.append(tr)
                while len(ctx) > context_size:
                    ctx.pop(0)
            else:
                ctx.clear()            
        except Exception as e:
            if cb: 
                #logger.error("Error translating segment %s", i)
                cb(i, total, f"Error segment {i+1}: {e}", False)
        if i < total - 1: time.sleep(delay)

    paths = []
    seen = set()

    target_suffix = '.xlf' if Path(inp).suffix.lower() == '.xlf' else '.xliff'

    for fmt in output_formats:
        if fmt in ('original', 'xliff'):
            out_path = str(Path(out).with_suffix(target_suffix))

            if out_path not in seen:
                tree.write(out_path, encoding='utf-8', xml_declaration=True)
                seen.add(out_path)
                paths.append(out_path)

        elif fmt == 'bilingual_xlsx':
            paths.append(generate_bilingual_xlsx(pairs, sl, tl, out + '_bilingual'))

        elif fmt == 'tmx':
            paths.append(generate_tmx(pairs, sl, tl, out + '_bilingual'))

        elif fmt == 'tsv':
            paths.append(generate_tsv(pairs, sl, tl, out + '_bilingual'))

    if cb: cb(total, total, f"XLIFF completed — {len(paths)} file(s) generated.", True)
    return paths

# ── Dispatcher ─────────────────────────────────────────────────────────────
def translate_file(input_path, output_path, api_key, source_lang, target_lang,
                   model, delay, skip_translated, style_instructions,
                   output_formats=None, progress_cb=None, stop_check_cb=None,
                   use_segmentation=False, context_size=3):
    """
    Dispatcher principal.
    
    Nouveaux paramètres:
        use_segmentation: Active la segmentation par phrase
        context_size: Taille du contexte (0, 3, ou 5 phrases)
    """
    if not output_formats:
        output_formats = ['original']
    
    ext = Path(input_path).suffix.lower()
    if ext not in SUPPORTED:
        raise ValueError(f"Unsupported format: {ext}")
    
    try:
        if ext == '.docx':
            return translate_docx(
                input_path, output_path, api_key, source_lang, target_lang,
                model, delay, style_instructions, output_formats, 
                progress_cb, stop_check_cb, use_segmentation, context_size
            )
        elif ext == '.xlsx':
            return translate_xlsx(
                input_path, output_path, api_key, source_lang, target_lang,
                model, delay, style_instructions, output_formats,
                progress_cb, stop_check_cb, use_segmentation, context_size
            )
        else:
            return translate_xliff_file(
                input_path, output_path, api_key, source_lang, target_lang,
                model, delay, skip_translated, style_instructions, output_formats,
                progress_cb, stop_check_cb, context_size
            )
    except TranslationStoppedException:
        return []